import asyncio
import io
import os
from typing import Optional
import boto3
from docx import Document
import fitz  # PyMuPDF
from PIL import Image


# Global clients for efficiency
_textract_client = None
_s3_client = None


def get_textract_client():
    """Lazy loading of AWS Textract client"""
    global _textract_client
    if not _textract_client:
        _textract_client = boto3.client('textract', region_name='us-east-1')
    return _textract_client


def get_s3_client():
    """Lazy loading of AWS S3 client"""
    global _s3_client
    if not _s3_client:
        _s3_client = boto3.client('s3', region_name='us-east-1')
    return _s3_client


async def extract_text(content: bytes, filename: str) -> str:
    """
    Smart text extraction based on file type and size
    Optimized for speed and accuracy
    """

    method = get_extraction_method(filename, len(content))
    print(f"Using extraction method: {method} for {filename}")

    try:
        if method == 'aws_textract':
            return await extract_with_textract(content)
        elif method == 'pymupdf':
            return extract_with_pymupdf(content)
        elif method == 'python_docx':
            return extract_with_docx(content)
        elif method == 'direct_read':
            return content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported extraction method: {method}")

    except Exception as e:
        print(f"Primary extraction failed with {method}: {str(e)}")
        # Fallback to AWS Textract for complex files
        if method != 'aws_textract':
            print("Attempting fallback to AWS Textract...")
            try:
                return await extract_with_textract(content)
            except Exception as fallback_error:
                print(f"Fallback extraction failed: {str(fallback_error)}")
                raise Exception(f"All extraction methods failed. Primary: {str(e)}, Fallback: {str(fallback_error)}")
        else:
            raise e


def get_extraction_method(filename: str, file_size: int) -> str:
    """
    Smart routing based on file type and characteristics
    Optimized for performance and accuracy
    """

    ext = filename.lower().split('.')[-1] if '.' in filename else ''

    # Images always use AWS Textract for OCR
    if ext in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']:
        return 'aws_textract'

    # PDFs: smart selection based on size
    elif ext == 'pdf':
        if file_size > 5_000_000:  # 5MB+ likely complex/scanned
            return 'aws_textract'
        else:
            return 'pymupdf'  # Try fast extraction first

    # Word documents
    elif ext in ['doc', 'docx']:
        return 'python_docx'

    # Plain text
    elif ext in ['txt', 'rtf']:
        return 'direct_read'

    # Unknown formats - try AWS Textract
    else:
        return 'aws_textract'


async def extract_with_textract(content: bytes) -> str:
    """
    Extract text using AWS Textract
    Best for: PDFs, Images, Complex documents
    """

    try:
        client = get_textract_client()

        # Textract has 10MB limit for direct API calls
        if len(content) > 10_000_000:
            raise ValueError("File too large for Textract direct processing (>10MB)")

        print(f"Processing {len(content)} bytes with AWS Textract...")

        response = await asyncio.to_thread(
            client.detect_document_text,
            Document={'Bytes': content}
        )

        # Extract text from Textract response
        text_blocks = []

        for block in response.get('Blocks', []):
            if block['BlockType'] == 'LINE':
                text_blocks.append(block.get('Text', ''))

        extracted_text = '\n'.join(text_blocks)

        print(f"AWS Textract extracted {len(extracted_text)} characters")
        return extracted_text

    except Exception as e:
        print(f"AWS Textract extraction failed: {str(e)}")
        raise e


def extract_with_pymupdf(content: bytes) -> str:
    """
    Extract text using PyMuPDF
    Best for: Simple PDFs, Fast processing
    """

    try:
        print(f"Processing {len(content)} bytes with PyMuPDF...")

        # Open PDF from bytes
        pdf_stream = io.BytesIO(content)
        pdf_doc = fitz.open(stream=pdf_stream, filetype="pdf")

        text_blocks = []

        for page_num in range(pdf_doc.page_count):
            page = pdf_doc[page_num]
            text = page.get_text()

            if text.strip():
                text_blocks.append(text)

        pdf_doc.close()
        extracted_text = '\n'.join(text_blocks)

        # Check if extraction was successful
        if len(extracted_text.strip()) < 50:
            raise ValueError("PyMuPDF extracted insufficient text - likely scanned PDF")

        print(f"PyMuPDF extracted {len(extracted_text)} characters")
        return extracted_text

    except Exception as e:
        print(f"PyMuPDF extraction failed: {str(e)}")
        raise e


def extract_with_docx(content: bytes) -> str:
    """
    Extract text using python-docx
    Best for: .docx and .doc files
    """

    try:
        print(f"Processing {len(content)} bytes with python-docx...")

        # Open document from bytes
        doc_stream = io.BytesIO(content)
        document = Document(doc_stream)

        text_blocks = []

        # Extract paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_blocks.append(paragraph.text)

        # Extract tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_blocks.append(' | '.join(row_text))

        extracted_text = '\n'.join(text_blocks)

        print(f"python-docx extracted {len(extracted_text)} characters")
        return extracted_text

    except Exception as e:
        print(f"python-docx extraction failed: {str(e)}")
        raise e


def validate_extracted_text(text: str, filename: str) -> str:
    """
    Validate and clean extracted text
    """

    if not text or not isinstance(text, str):
        raise ValueError(f"No text extracted from {filename}")

    # Clean the text
    text = text.strip()

    # Check minimum length
    if len(text) < 50:
        raise ValueError(f"Insufficient text extracted from {filename} ({len(text)} characters)")

    # Basic content validation
    if not any(char.isalpha() for char in text):
        raise ValueError(f"No readable text found in {filename}")

    print(f"Text validation passed: {len(text)} characters extracted")
    return text


def estimate_processing_time(filename: str, file_size: int) -> dict:
    """
    Estimate processing time based on file characteristics
    """

    method = get_extraction_method(filename, file_size)

    time_estimates = {
        'aws_textract': {'min': 6, 'max': 12, 'avg': 9},
        'pymupdf': {'min': 1, 'max': 3, 'avg': 2},
        'python_docx': {'min': 1, 'max': 2, 'avg': 1.5},
        'direct_read': {'min': 0.1, 'max': 0.5, 'avg': 0.2}
    }

    return {
        'method': method,
        'estimated_seconds': time_estimates.get(method, {'avg': 5})['avg'],
        'file_size_mb': round(file_size / 1024 / 1024, 2)
    }